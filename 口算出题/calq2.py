import configparser
from docx import Document
from docx.shared import Inches, Pt
from random import randint, choice
from datetime import datetime

def safe_operation(lower, upper, operation):
    """ Helper function to generate safe operations within a given range """
    while True:
        a = randint(lower, upper)
        b = randint(lower, upper)
        result = a + b if operation == '+' else a - b
        if lower <= result <= upper:
            return a, b, result

def safe_operation2(lower, upper):
    """ Helper function to generate safe operations within a given range """
    while True:
        a = randint(lower, upper)
        b = randint(lower, upper)
        return a, b

def safe_operation3(lower, upper):
    """ Helper function to generate safe operations within a given range """

    a = randint(lower, upper)
    b = randint(lower, upper)
    operation = choice(['+', '-'])
    if operation == '+':
        result = a + b
        if lower <= result <= upper:
            return a, b, result,operation

    operation = '-'
    result = a - b
    if lower <= result <= upper:
        return a, b, result,operation
    else:
        result = b - a
        if lower <= result <= upper:
           return b, a, result,operation
        else:
            return 0, 0, 0,operation


def generate_problem(problem_type):
    if problem_type == 0:
        # 数据在10以内,得数大于0的加减法
        a, b, result = safe_operation(1, 10, choice(['+', '-']))
        return f"{a} + {b} = " if result == a + b else f"{a} - {b} = "
    elif problem_type == 1:
        # 数据在10以内,得数大于0的的3个数据的加减法
        a, b, _ = safe_operation(1, 10, choice(['+', '-']))
        c, _, result = safe_operation(1, 10, choice(['+', '-']))
        operations = [choice(['+', '-']) for _ in range(2)]
        return f"{a} {operations[0]} {b} {operations[1]} {c} = "
    elif problem_type == 2:
        # 数据在20以内,得数大于0的加减法
        a, b, result,operation = safe_operation3(1, 20)
        return f"{a} {operation} {b} = "
    elif problem_type == 3:
        # 得数在0到100之间加减法
        a, b, result = safe_operation(0, 100, choice(['+', '-']))
        return f"{a} + {b} = " if result == a + b else f"{a} - {b} = "
    elif problem_type == 4:
        # 得数在0到100之间的3个数据的连加减法
        a, b, _ = safe_operation(0, 100, choice(['+', '-']))
        c, _, result = safe_operation(0, 100, choice(['+', '-']))
        operations = [choice(['+', '-']) for _ in range(2)]
        return f"{a} {operations[0]} {b} {operations[1]} {c} = "
    elif problem_type == 5:
        # 得数在0到200之间的3个数据的连加减法
        a, b, _ = safe_operation(0, 200, choice(['+', '-']))
        c, _, result = safe_operation(0, 200, choice(['+', '-']))
        operations = [choice(['+', '-']) for _ in range(2)]
        return f"{a} {operations[0]} {b} {operations[1]} {c} = "
    elif problem_type == 6:
        # 数据在1000以内,得数大于0的加减法
        a, b, result = safe_operation(1, 1000, choice(['+', '-']))
        return f"{a} + {b} = " if result == a + b else f"{a} - {b} = "
    elif problem_type == 7:
        # 数据在1000以内,得数大于0的加减法
        a, b = safe_operation2(100, 1000)
        operations = [choice(['+', '-']) for _ in range(1)]
        return f"{a} {operations[0]} {b} = "
# 读取配置文件
config = configparser.ConfigParser()
config.read('config.ini')
num_pages = int(config['SETTINGS']['num_pages'])
problem_type = int(config['SETTINGS']['problem_type'])

# 创建Word文档
doc = Document()
for page_number in range(num_pages):
    # 添加日期、时间和成绩行
    date_time_score = doc.add_paragraph("日期：                           计算时间：                             成绩：                             ")
    date_time_score.style = doc.styles['Normal']
    for run in date_time_score.runs:
        run.font.size = Pt(14)  # 字体大小设为14磅

    # 添加表格
    table = doc.add_table(rows=17, cols=3)
    table.autofit = False
    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            cell.text = generate_problem(problem_type)
            cell.width = Inches(2.3)  # 设置每个单元格的宽度以均匀分布题目
            cell.paragraphs[0].runs[0].font.size = Pt(14)  # 设置表格内字体大小为14磅
    # 仅在不是最后一页时添加新的章节
    if page_number < num_pages - 1:
        doc.add_section()

# 动态生成文件名
current_time = datetime.now().strftime("%Y%m%d_%H%M%S")
filename = f"{problem_type}_{current_time}.docx"

# 保存文档
doc.save(filename)
